# prism/backend/app/utils/file_parser.py
"""Utilities for extracting text from supported knowledge sources."""

from pathlib import Path


def extract_text(file_path: str) -> str:
    """Extract plain text from a supported file."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext == ".doc":
        raise ValueError(".doc parsing is not available; convert to .docx or upload for metadata only")
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext == ".xlsx":
        return _extract_xlsx(file_path)
    if ext in (".md", ".txt", ".markdown"):
        return Path(file_path).read_text(encoding="utf-8")
    raise ValueError(f"不支持的文件类型: {ext}")


def _extract_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF

    with fitz.open(file_path) as doc:
        text_parts = [f"[[PAGE:{index}]]\n{page.get_text()}" for index, page in enumerate(doc, start=1)]
    return "\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)


def _extract_xlsx(file_path: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    try:
        text_parts = []
        for sheet in wb.worksheets:
            text_parts.append(f"## Worksheet {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(cells):
                    text_parts.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(text_parts)


def count_pages(file_path: str) -> int | None:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        import fitz

        with fitz.open(file_path) as doc:
            return len(doc)
    if ext == ".docx":
        from docx import Document

        doc = Document(file_path)
        return max(1, len(doc.paragraphs))
    if ext in (".md", ".txt", ".markdown"):
        return None
    return None


def extract_url(url: str) -> str:
    """Fetch a web page and strip it to readable text."""
    import httpx
    import re

    resp = httpx.get(url, timeout=30, follow_redirects=True, headers={"User-Agent": "Prism/1.0"})
    resp.raise_for_status()
    html = resp.text
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
